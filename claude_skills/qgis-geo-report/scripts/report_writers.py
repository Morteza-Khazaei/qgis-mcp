"""Report-assembly primitives for qgis-geo-report (HTML / Markdown / Word).

Copy into the report's codes/ folder. The builders compose content (captions,
takeaways, tables) — these functions guarantee the visual system stays identical
across projects. Same section content goes to all three formats.
"""
import base64

# ------------------------------------------------------------------ HTML
CSS = """
 body{font-family:system-ui,-apple-system,"Segoe UI",sans-serif;color:#0b0b0b;background:#f9f9f7;
      margin:0;line-height:1.55}
 main{max-width:880px;margin:0 auto;padding:32px 24px 64px}
 h1{font-size:1.7em;margin:.2em 0 .1em} h2{font-size:1.25em;margin-top:2em;border-bottom:1px solid #e1e0d9;padding-bottom:.25em}
 .sub{color:#52514e;margin-bottom:1.5em}
 .card{background:#fcfcfb;border:1px solid rgba(11,11,11,.10);border-radius:8px;padding:18px 20px;margin:14px 0}
 img{max-width:100%;height:auto;border-radius:4px}
 figcaption{font-size:.86em;color:#52514e;margin-top:6px}
 figure{margin:0}
 table{border-collapse:collapse;width:100%;font-size:.9em}
 th,td{text-align:left;padding:6px 10px;border-bottom:1px solid #e1e0d9}
 th{color:#52514e;font-weight:600} td.num{font-variant-numeric:tabular-nums}
 .tag{font-size:.75em;padding:1px 7px;border-radius:9px;font-weight:600}
 .ok{background:#e2f2e2;color:#006300} .asm{background:#fdf0d7;color:#8a5a00}
 .take{background:#eef4fc;border-left:3px solid #2a78d6;padding:10px 14px;border-radius:0 6px 6px 0;margin:12px 0}
 .note{color:#898781;font-size:.95em}
 .hero{font-size:2.4em;font-weight:700;color:#2a78d6;line-height:1.1}
 .heror{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin:16px 0}
 .hero small{display:block;font-size:.38em;font-weight:400;color:#52514e;margin-top:4px}
"""


def b64(figures_dir, name):
    with open(f"{figures_dir}/{name}.png", "rb") as f:
        return "data:image/png;base64," + base64.b64encode(f.read()).decode()


def html_page(title, sub, body):
    return (f'<meta charset="utf-8">\n<title>{title}</title>\n<style>{CSS}</style>\n'
            f'<main>\n<h1>{title}</h1>\n<div class="sub">{sub}</div>\n{body}\n</main>')


def html_fig(figures_dir, name, caption):
    return (f'<figure class="card"><img src="{b64(figures_dir, name)}">'
            f'<figcaption>{caption}</figcaption></figure>')


def html_take(text):
    return f'<div class="take"><b>Takeaway.</b> {text}</div>'


def html_table(header, rows, note="", bold_last=False):
    h = "<tr>" + "".join(f"<th>{c}</th>" for c in header) + "</tr>"
    body = ""
    for i, r in enumerate(rows):
        b = bold_last and i == len(rows) - 1
        body += "<tr>" + "".join(
            f'<td class="num">{"<b>%s</b>" % c if b else c}</td>' if j else
            f'<td>{"<b>%s</b>" % c if b else c}</td>'
            for j, c in enumerate(r)) + "</tr>"
    n = f'<p class="note">{note}</p>' if note else ""
    return f'<div class="card"><table>{h}{body}</table>{n}</div>'


def html_heroes(items):
    """items: list of (big_value, small_caption)."""
    cells = "".join(f'<div class="card"><span class="hero">{v}<small>{c}</small></span></div>'
                    for v, c in items)
    return f'<div class="heror">{cells}</div>'


# ------------------------------------------------------------------ Markdown
def md_table(header, rows):
    out = ["| " + " | ".join(header) + " |", "|" + "---|" * len(header)]
    out += ["| " + " | ".join(str(c) for c in r) + " |" for r in rows]
    return "\n".join(out)


def md_fig(name, caption):
    return f"![{name}](figures/{name}.png)\n*{caption}*\n"


# ------------------------------------------------------------------ Word
def make_doc():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(10.5)
    return doc


def docx_fig(doc, figures_dir, name, caption):
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc.add_picture(f"{figures_dir}/{name}.png", width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption)
    p.runs[0].italic = True; p.runs[0].font.size = Pt(9)


def docx_table(doc, header, rows, note="", bold_last=False):
    from docx.shared import Pt
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(header):
        t.rows[0].cells[j].text = str(c)
        t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, r in enumerate(rows, start=1):
        for j, c in enumerate(r):
            t.rows[i].cells[j].text = str(c)
            if bold_last and i == len(rows):
                t.rows[i].cells[j].paragraphs[0].runs[0].bold = True
    if note:
        p = doc.add_paragraph(note)
        p.runs[0].italic = True; p.runs[0].font.size = Pt(8.5)


def docx_take(doc, text):
    p = doc.add_paragraph()
    p.add_run("Takeaway: " + text).bold = True


def docx_heroes(doc, items):
    from docx.shared import Pt, RGBColor
    for v, c in items:
        p = doc.add_paragraph()
        r = p.add_run(str(v) + "  ")
        r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x2A, 0x78, 0xD6)
        p.add_run(c).font.size = Pt(9)
